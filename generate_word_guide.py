import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def create_guide():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("PANDUAN LENGKAP PRESENTASI KE-2\n& CARA MEMBACA ERD BRIGHTDOR")
    t_run.font.name = "Arial"
    t_run.font.size = Pt(20)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(142, 27, 55) # Maroon

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("BrightDor Premier Wedding Marketplace — Progress 10 September 2026\nDisusun Khusus Sebagai Contekan & Panduan Menghadapi Dosen Penguji")
    s_run.font.name = "Arial"
    s_run.font.size = Pt(11)
    s_run.font.italic = True
    s_run.font.color.rgb = RGBColor(90, 80, 84)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for headings
    def add_custom_heading(text, level=1):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.bold = True
        if level == 1:
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(142, 27, 55)
            # Underline bar
        else:
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(168, 134, 71) # Gold
        return h

    # Section 1: Ringkasan Status Proyek
    add_custom_heading("1. STATUS PROYEK & KESIAPAN SISTEM", 1)
    p = doc.add_paragraph()
    p.add_run("• Status Utama: ").bold = True
    p.add_run("Keseluruhan Pengguna SUDAH JALAN (Couple, Vendor, Admin saling terintegrasi).\n")
    p.add_run("• Tech Stack: ").bold = True
    p.add_run("Laravel 11, Filament v5 (Panel Admin & Vendor), Tailwind CSS, MySQL, Spatie Media Library.\n")
    p.add_run("• Kepatuhan Pengujian: ").bold = True
    p.add_run("54 Automated Tests (174 assertions) 100% Passed (Hijau).\n")
    p.add_run("• Halaman Presentasi Interaktif: ").bold = True
    p.add_run("Tersedia di http://127.0.0.1:8000/presentasi/index.html (statis, mandiri, bebas lag).\n")

    # Section 2: Penjelasan 3 Peran Pengguna
    add_custom_heading("2. PENJELASAN 3 PERAN PENGGUNA (USER ROLES)", 1)
    doc.add_paragraph(
        "Aplikasi menerapkan Role-Based Access Control (RBAC) ketat via Spatie Permission yang memisahkan hak akses ke dalam 3 jenis user:"
    )

    table_roles = doc.add_table(rows=1, cols=3)
    table_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_roles.rows[0].cells
    hdr_cells[0].text = "Peran (Role)"
    hdr_cells[1].text = "Hak Akses & Dashboard"
    hdr_cells[2].text = "Fungsi Utama dalam Alur Bisnis"

    for cell in hdr_cells:
        set_cell_background(cell, "8E1B37")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                r.font.size = Pt(10)

    roles_data = [
        ("Couple (Customer)", "Frontend Publik & Dashboard Booking Saya (/booking-saya)", "Mencari vendor, filter kategori/kota, booking tanggal, ajukan penawaran nego budget, memantau status pesanan, memberi rating review resmi, dan lupa password mandiri."),
        ("Vendor (Mitra Usaha)", "Panel Mandiri Filament (/vendor/login & /vendor/dashboard)", "Mengelola profil usaha, foto portofolio, katalog paket jasa (services), menyetujui pesanan masuk, mengunggah dokumen legalitas, dan mengajukan payout saldo."),
        ("Admin (Superuser)", "Panel Pusat Filament (/admin/login & /admin/dashboard)", "Pengawas operasional: validasi pembayaran semi-otomatis, kurasi vendor verified & featured, pengawasan arus kas transaksi & komisi, serta audit keamanan.")
    ]

    for role, dash, desc in roles_data:
        row_cells = table_roles.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = dash
        row_cells[2].text = desc
        for i, c in enumerate(row_cells):
            set_cell_background(c, "FDF5F7" if i == 0 else "FFFFFF")
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    # Section 3: 3 Tipe / Pilar Admin
    add_custom_heading("3. PENJELASAN 3 PILAR ARSITEKTUR DASHBOARD ADMIN", 1)
    doc.add_paragraph(
        "Sesuai arahan presentasi ke-2, jika dosen meminta penjelasan mendalam mengenai peran Admin, gunakan 3 poin pilar ini:"
    )

    p1 = doc.add_paragraph()
    p1.add_run("A. Validasi Pembayaran Semi-Auto:\n").bold = True
    p1.add_run(
        "Transaksi dengan payment gateway (Midtrans/Xendit) tervalidasi otomatis oleh webhook callback. "
        "Namun untuk metode pembayaran manual (transfer langsung), admin memiliki tombol verifikasi satu-klik "
        "untuk mencocokkan mutasi bank sebelum mengubah status transaksi menjadi berhasil."
    )

    p2 = doc.add_paragraph()
    p2.add_run("B. Manajemen Menu Terstruktur (Admin Sibuk saat Banyak Menu):\n").bold = True
    p2.add_run(
        "Meskipun admin mengelola banyak modul sekaligus (Users, Vendors, Bookings, Transactions, Testimonials, CMS), "
        "antarmuka Filament mengelompokkannya secara rapi ke dalam klaster navigasi (Master Data, Marketplace, Keuangan, CMS) "
        "dilengkapi fitur pencarian global (Ctrl+K) agar alur kerja admin tetap cepat dan tidak membingungkan."
    )

    p3 = doc.add_paragraph()
    p3.add_run("C. Otomatisasi Sistem (Sudah Tervalidasi Sendiri):\n").bold = True
    p3.add_run(
        "Platform dirancang tidak bergantung pada campur tangan manual untuk hal-hal berulang: "
        "1) Komisi platform dihitung otomatis dari transaksi lunas; "
        "2) Rating rata-rata vendor otomatis diperbarui saat ulasan baru masuk; "
        "3) Status langganan vendor (active vs expired) diperbarui otomatis berdasarkan timestamp masa berlaku."
    )

    # Section 4: PANDUAN CARA MEMBACA ERD (Paling Penting!)
    add_custom_heading("4. PANDUAN LANGKAH DEMI LANGKAH CARA MEMBACA ERD", 1)
    doc.add_paragraph(
        "JANGAN membaca tabel satu per satu dari pojok kiri ke kanan! Dosen ingin mendengar alur bisnis yang logis. "
        "Gunakan alur 5 langkah di bawah ini sambil mengarahkan mouse:"
    )

    steps = [
        ("Langkah 1: Membuka Presentasi ERD", 
         "\"Bapak/Ibu Dosen, ini adalah skema ERD BrightDor yang terdiri dari 37 tabel relasional dalam 7 modul. Diagram ini didesain ternormalisasi, menerapkan prinsip One-User-One-Vendor, relasi polimorfik, dan jejak audit forensik.\""),
        
        ("Langkah 2: Tunjuk Tabel 'users' & 'roles'", 
         "\"Titik awal sistem bermula dari tabel users. Tabel ini menampung semua akun (Admin, Vendor, Couple). Khusus vendor, tabel users menyimpan kolom vendor_subscription_status dan vendor_subscription_expires_at untuk mengontrol lisensi masa aktif akun vendor.\""),
        
        ("Langkah 3: Tunjuk Relasi 'users' → 'vendors' → 'services'", 
         "\"Ketika user mendaftar sebagai vendor, dibuat relasi 1:1 (Unique) antara users.id dengan vendors.user_id. Satu user hanya punya satu vendor. Vendor dikelompokkan ke vendor_categories, dan memiliki relasi 1:N ke tabel services untuk paket harga, durasi, dan kapasitas tamu.\""),
        
        ("Langkah 4: Tunjuk Relasi Inti 'bookings' → 'reviews'", 
         "\"Tabel bookings adalah persimpangan transaksi utama: mempertemukan user_id (couple), vendor_id, dan service_id. Tabel ini mencatat tanggal acara dan status pesanan (pending → confirmed → on_progress → completed). Setelah status completed, barulah couple bisa mengisi tabel reviews. Relasi booking_id dibuat unique di reviews agar tidak bisa spam ulasan palsu.\""),
        
        ("Langkah 5: Tunjuk Relasi Polimorfik 'transactions' & 'payouts'", 
         "\"Di modul keuangan, tabel transactions menggunakan relasi polimorfik (payable_type dan payable_id). Ini memungkinkan satu tabel transaksi memproses pembayaran booking vendor maupun pembelian undangan digital secara efisien. Sisa saldo bersih vendor kemudian dicairkan melalui tabel payouts yang diproses oleh admin.\""),
        
        ("Langkah 6: Tunjuk Modul Media & Keamanan", 
         "\"Aset foto portofolio dikelola secara polimorfik oleh Spatie Media Library di tabel media tanpa membebani tabel bisnis. Seluruh log perubahan data penting dicatat di audit_logs untuk forensik keamanan, dan password pengguna diamankan secara mandiri melalui password_reset_tokens.\"")
    ]

    for title_s, script_s in steps:
        p_step = doc.add_paragraph()
        p_step.paragraph_format.left_indent = Inches(0.2)
        p_step.add_run(f"▶ {title_s}\n").bold = True
        run_script = p_step.add_run(f"Kalimat yang diucapkan: {script_s}\n")
        run_script.italic = True
        run_script.font.color.rgb = RGBColor(40, 34, 36)

    # Section 5: Security & Hosting
    add_custom_heading("5. KEAMANAN & STRATEGI HOSTING (SESUAI ARAHAN TANGGAL 10)", 1)
    doc.add_paragraph(
        "Bahan presentasi untuk menjawab pertanyaan dosen terkait kesiapan server dan keamanan data:"
    )

    sec_points = [
        ("Enkripsi Bcrypt Password: ", "Password tidak pernah disimpan dalam bentuk plaintext, melainkan di-hash satu arah dengan Bcrypt (work factor 12)."),
        ("Lupa Password Mandiri (Self-Service): ", "Admin dilarang mengubah password user secara langsung untuk menutup celah rekayasa sosial (social engineering). User mereset sendiri via email token."),
        ("Proteksi CSRF & XSS: ", "Setiap request form dilindungi token CSRF dan rendering Blade menerapkan HTML escaping otomatis."),
        ("Efisiensi Database Hosting: ", "Sesuai arahan: hosting database disesuaikan standar kebutuhan (indexing tepat, normalisasi bersih) sehingga ringan di awal tanpa membebani biaya sewa server."),
        ("Skalabilitas Domain & Hosting: ", "Sesuai arahan: saat domain atau shared hosting awal mulai lemah akibat lonjakan trafik, sistem siap di-upgrade ke VPS atau cloud container (DigitalOcean/AWS) secara seamless.")
    ]

    for k, v in sec_points:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.left_indent = Inches(0.2)
        p_sec.add_run("• " + k).bold = True
        p_sec.add_run(v)

    # Section 6: Q&A Cheat Sheet
    add_custom_heading("6. CONTEKAN TANYA JAWAB (Q&A CHEAT SHEET)", 1)
    
    qa_list = [
        ("Kenapa tabel transactions dan media memakai relasi polimorfik?",
         "Karena relasi polimorfik memungkinkan satu tabel menangani banyak model berbeda. Tabel transactions bisa memproses booking vendor dan pesanan undangan digital sekaligus tanpa membuat dua tabel terpisah. Ini membuat struktur bersih dan mengikuti prinsip DRY (Don't Repeat Yourself)."),
        
        ("Bagaimana mencegah vendor membuat ulasan palsu sendiri?",
         "Tabel reviews memiliki foreign key unique ke booking_id. Di aplikasi divalidasi bahwa ulasan hanya dapat dibuat jika status booking sudah 'completed' dan user_id pembuat ulasan adalah couple yang benar-benar memesan."),
        
        ("Kenapa ada 37 tabel? Apakah tidak terlalu banyak?",
         "Jumlah 37 tabel ini mencerminkan arsitektur sistem enterprise yang matang: 8 tabel autentikasi/RBAC, 4 tabel vendor, 2 tabel booking/review, 5 tabel undangan digital, 3 tabel keuangan, 5 tabel CMS, dan 10 tabel penunjang sistem (queue, audit log, cache, media). Setiap tabel memiliki single-responsibility sehingga performa database tetap optimal."),
        
        ("Apa peran admin jika validasi pembayaran sudah semi-auto?",
         "Admin tetap memegang kendali atas persetujuan mutasi transfer manual, pencairan dana vendor (payout), moderasi dokumen verifikasi mitra (KTP/NIB), dan memantau audit log jika ada anomali aktivitas pengguna.")
    ]

    for q, a in qa_list:
        p_qa = doc.add_paragraph()
        p_qa.paragraph_format.left_indent = Inches(0.2)
        p_qa.add_run(f"Tanya: \"{q}\"\n").bold = True
        r_ans = p_qa.add_run(f"Jawab: \"{a}\"\n")
        r_ans.font.color.rgb = RGBColor(14, 98, 59) # Emerald Green

    # Output path
    out_dir = r"C:\Users\ADVAN\brightdor\docs"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Panduan_Presentasi_ERD_dan_Fitur_BrightDor.docx")
    doc.save(out_path)
    print(f"Word document saved to: {out_path}")

create_guide()
